from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # 段落对齐
from docx.shared import Inches, Pt  # 段落缩进
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
import os
import datetime
import time
import requests
from lxml import etree

session = requests.Session()
first_url = 'http://www.nmc.cn/publish/observations/environmental.html'

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.111 Safari/537.36'
}


def real():
    response = session.get(url=first_url, headers=headers)
    response.encoding = "utf-8"
    res = response.text
    html = etree.HTML(res)
    title = html.xpath('//div[@class="writing"]/div')
    ret = html.xpath('//div[@class="writing"]/p')
    image = html.xpath('//div[@class="writing"]/div/img/@src')
    text_name = title + ret
    text_x = []
    for i in text_name:
        if i.text != None:
            text_x.append(i.text)

    for i in range(len(image)):
        if image[i] != None:
            res = session.get(image[i], headers=headers)
            with open('county_image/' + str(i) + '.jpg', 'wb') as f:
                f.write(res.content)
    return text_x


def headings1(document, text):
    # 一级标题样式设计
    run = document.add_heading('', level=1)
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runs = run.add_run(text)
    runs.font.size = Pt(22)
    runs.bold = True
    runs.font.color.rgb = RGBColor(0, 0, 0)
    runs.font.name = u'Times New Roman'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def zw(document, text):
    # 正文字体样式
    document.styles['Normal'].font.name = u'Times New Roman'  # 正文样式
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 正文样式设计
    paragraph = document.add_paragraph()
    paragraph.add_run(text).font.size = Pt(12)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.first_line_indent = Inches(0.25)  # 首行缩进
    paragraph_format.space_before = Pt(2)  # 段前段后间距设置
    paragraph_format.space_after = Pt(2)
    paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
    paragraph_format.line_spacing = Pt(18)  # 固定值18磅
    paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 多倍行距
    paragraph_format.line_spacing = 1.25


def pic_tab(document, text):
    paragraph = document.add_paragraph()
    paragraph.add_run(text).font.size = Pt(10.5)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_before = Pt(2)  # 段前段后间距设置
    paragraph_format.space_after = Pt(2)


def pics(path, document, table_style, size):
    dir_list = os.listdir(path)
    n = 1
    for k in dir_list:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(path + '/' + k, width=Inches(size))
        pic_tab(document, '图{0} 全国预报图'.format(n))
        n += 1


def hed_fot(document, headtext, footext):
    section = document.sections[0]
    section.header_distance = Cm(1.9)
    section.footer_distance = Cm(1.9)
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    paragraph.add_run(headtext).font.size = Pt(10.5)
    paragraph.add_run().font.underline = True
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format.space_before = Pt(0)  # 上行间距
    paragraph_format.space_after = Pt(0)  # 下行间距
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    paragraph.add_run(footext).font.size = Pt(10.5)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


# 每发完一次删除文件夹下的图片
def del_files(path_file):
    ls = os.listdir(path_file)
    for i in ls:
        f_path = os.path.join(path_file, i)
        # 判断是否是一个目录,若是,则递归删除
        if os.path.isdir(f_path):
            del_files(f_path)
        else:
            os.remove(f_path)


while True:
    try:
        document = Document()
        headings1(document, "环境气象公报")
        tex = real()
        for i in tex:
            zw(document, i)
        document.add_page_break()
        table_style = 'Table Grid'
        path = r'county_image'
        pics(path, document, table_style, 4)
        headtext = '中央气象台环境气象公报'
        footext = '佳华科技生态环境研究院'
        hed_fot(document, headtext, footext)
        now_time = datetime.datetime.strftime(datetime.datetime.now(), '%Y-%m-%d %H')
        #         if int(now_time[-2:])>12:
        #             tet = '下午'
        #         else:
        #             tet = '上午'
        document.save('file/{0}.docx'.format(now_time))
        del_files(path)
        print('完成')
    #     time.sleep(43200)
    except:
        print('NO')
        time.sleep(60)