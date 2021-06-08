from docx import Document


document = Document()
# document = Document('existing-document-file.docx')

document.add_heading('The role of dolphins', level=2)  #设置标题

document.add_page_break()   #新增分页符

table = document.add_table(rows=2, cols=2) #创建表格
cell = table.cell(0, 1)
cell.text = '这是第一行第二列的单元格'    #对单元格进行赋值
row = table.add_row()  #表格新增行
table.style = 'LightShading-Accent1'  #设置表格样式

document.add_picture('data/demo.png')   #插入图片

from docx.shared import Inches

document.add_picture('data/demo.png', width=Inches(1.0), height=Inches(1.0))   #调整图片大小

paragraph = document.add_paragraph('这是一个样式为 ListBullet 的段落')
paragraph.style = 'List Bullet'        #段落操作


from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING   #段落对齐

# LEFT      =>  左对齐
# CENTER    =>  文字居中
# RIGHT     =>  右对齐
# JUSTIFY   =>  文本两端对齐

paragraph = document.add_paragraph("你说啥")
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



from docx.shared import Inches,Pt            #段落缩进

paragraph = document.add_paragraph("你说啥")
paragraph_format = paragraph.paragraph_format
paragraph_format.left_indent = Inches(0.5)


paragraph_format.first_line_indent = Inches(-0.25)  #首行缩进


paragraph_format.space_before = Pt(18)  #段前段后间距设置
paragraph_format.space_after = Pt(12)

from docx.shared import Length

#SINGLE         =>  单倍行距（默认）
#ONE_POINT_FIVE =>  1.5倍行距
#DOUBLE2        =>  倍行距
#AT_LEAST       =>  最小值
#EXACTLY        =>  固定值
#MULTIPLE       =>  多倍行距

paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY  #固定值
paragraph_format.line_spacing = Pt(18)                # 固定值18磅
paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  #多倍行距
paragraph_format.line_spacing = 1.75                   # 1.75倍行间距


#widow_control      =>  孤行控制
#keep_with_next     =>  与下段同页
#page_break_before  =>  段前分页
#keep_together      =>  段中不分页

paragraph_format.keep_with_next = True


paragraph = document.add_paragraph() #设置字体
paragraph.add_run('这是一个带有')
paragraph.add_run('粗体').bold = True
paragraph.add_run('和')
paragraph.add_run('斜体').italic = True
paragraph.add_run('的段落。')




from docx.shared import RGBColor,Pt    #设置字体属性

#all_caps       =>  全部大写字母
#bold           =>  加粗
#color          =>  字体颜色
#complex_script =>  是否为“复杂代码”
#cs_bold        =>  “复杂代码”加粗
#cs_italic      =>  “复杂代码”斜体
#double_strike  =>  双删除线
#emboss         =>  文本以凸出页面的方式出现
#hidden         =>  隐藏
#imprint        =>  印记
#italic         =>  斜体
#name           =>  字体
#no_proof       =>  不验证语法错误
#outline        =>  显示字符的轮廓
#shadow         =>  阴影
#small_caps     =>  小型大写字母
#snap_to_grid   =>  定义文档网格时对齐网络
#strike         =>  删除线
#subscript      =>  下标
#superscript    =>  上标
#underline      =>  下划线

paragraph = document.add_paragraph()
paragraph.add_run('这是一个带有')
paragraph.add_run('颜色').font.color.rgb = RGBColor(54, 95, 145)
paragraph.add_run('的')
paragraph.add_run('大字').font.size = Pt(36)  # 字体大小设置，和word里面的字号相对应



# 自定义样式 Emphasis

paragraph = document.add_paragraph('这是一个带有')
paragraph.add_run('自定义样式', 'Emphasis')
paragraph.add_run('的段落')






document.save('file/new-file-name.docx')