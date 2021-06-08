from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt,Cm            #段落缩进




def hed_fot(document,headtext,footext):
    section = document.sections[0]
    section.header_distance = Cm(1.9)
    section.footer_distance = Cm(1.9)
    header = section.header
    paragraph=header.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    paragraph.add_run(headtext).font.size = Pt(10.5)
    paragraph.add_run().font.underline = True
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format.space_before = Pt(0)  # 上行间距
    paragraph_format.space_after = Pt(0)  # 下行间距
    footer = section.footer
    paragraph=footer.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    paragraph.add_run(footext).font.size = Pt(10.5)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
