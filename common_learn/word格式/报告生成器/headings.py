from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor


def headings1(document,text):
    #一级标题样式设计
    run = document.add_heading('',level=1)                    #标题样式
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runs = run.add_run(text)
    runs.font.size = Pt(22)
    runs.bold = True
    runs.font.color.rgb = RGBColor(0, 0, 0)
    runs.font.name=u'Times New Roman'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def headings2(document,text):
    #一级标题样式设计
    run = document.add_heading('',level=2)                    #标题样式
    run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    runs = run.add_run(text)
    runs.font.size = Pt(18)
    runs.bold = True
    runs.font.color.rgb = RGBColor(0, 0, 0)
    runs.font.name=u'Times New Roman'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def headings3(document,text):
    #一级标题样式设计
    run = document.add_heading('',level=3)                    #标题样式
    run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    runs = run.add_run(text)
    runs.font.size = Pt(15)
    runs.bold = True
    runs.font.color.rgb = RGBColor(0, 0, 0)
    runs.font.name=u'Times New Roman'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def headings4(document,text):
    #一级标题样式设计
    run = document.add_heading('',level=14)                    #标题样式
    run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    runs = run.add_run(text)
    runs.font.size = Pt(12)
    runs.bold = True
    runs.font.color.rgb = RGBColor(0, 0, 0)
    runs.font.name=u'Times New Roman'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def headings5(document,text):
    #一级标题样式设计
    run = document.add_heading('',level=12)                    #标题样式
    run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    runs = run.add_run(text)
    runs.font.size = Pt(12)
    runs.bold = True
    runs.font.color.rgb = RGBColor(0, 0, 0)
    runs.font.name=u'Times New Roman'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')