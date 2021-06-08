from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os


def pics(path,document,table_style,num,size):
    dir_list1 = os.listdir(path)
    if len(dir_list1) % num !=0:
        l=1
    else:
        l=0
    table = document.add_table(rows=(len(dir_list1)//num)+l, cols=num)

    print(dir_list1)
    n = 0
    for row in table.rows:
          for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    try:
                        run = paragraph.add_run()
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.add_picture(path+'\{}'.format(dir_list1[n]),width=Inches(size))
                        print(n)
                        n+=1
                    except:
                        n+=1


    table.style = table_style  #设置表格样式
#插入图片样式设计
